#!/usr/bin/env python3
"""
Script para generar documento Word con todas las figuras de análisis climático.
Adaptado para usar la configuración centralizada.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from organized.config import settings


FILE_STRUCTURE = {
    "1_series_temporales_temperatura": {
        "title": "1. Series Temporales de Temperatura (1980-2100)",
        "files": [
            ("01_Series_Temporales_Temperatura/temperatura_media_anual.png", "Temperatura media anual (°C)"),
            ("01_Series_Temporales_Temperatura/temperatura_maxima_anual.png", "Temperatura máxima anual (°C)"),
            ("01_Series_Temporales_Temperatura/temperatura_minima_anual.png", "Temperatura mínima anual (°C)"),
            ("01_Series_Temporales_Temperatura/warming_stripes_anomalias.png", "Warming Stripes - Anomalías de temperatura por escenario"),
        ]
    },
    "2_series_temporales_hidro": {
        "title": "2. Series Temporales de Variables Hidrológicas (1980-2100)",
        "files": [
            ("02_Series_Temporales_Hidrologicas/precipitacion_anual.png", "Precipitación anual (mm/año)"),
            ("02_Series_Temporales_Hidrologicas/evapotranspiracion_anual.png", "Evapotranspiración potencial anual (mm/año)"),
            ("02_Series_Temporales_Hidrologicas/balance_hidrico_anual.png", "Balance hídrico anual (mm/año)"),
        ]
    },
    "3_indicadores_sequia": {
        "title": "3. Indicadores de Sequía",
        "files": [
            ("03_Indicadores_Sequia/indice_aridez_serie_temporal.png", "Índice de Aridez (AI = P/PET) - Serie temporal y anomalías"),
            ("03_Indicadores_Sequia/dias_secos_consecutivos_serie_temporal.png", "Días Secos Consecutivos (CDD) - Serie temporal y anomalías"),
        ]
    },
    "4_climatologia_mensual_comparativa": {
        "title": "4. Climatología Mensual - Comparación Base vs Futuro",
        "files": [
            ("04_Climatologia_Mensual_Comparativa/ciclo_anual_precipitacion.png", "Ciclo anual de Precipitación (mm/mes)"),
            ("04_Climatologia_Mensual_Comparativa/ciclo_anual_evapotranspiracion.png", "Ciclo anual de Evapotranspiración (mm/mes)"),
            ("04_Climatologia_Mensual_Comparativa/ciclo_anual_balance_hidrico.png", "Ciclo anual de Balance Hídrico (mm/mes)"),
        ]
    },
    "5_climatologia_mensual_ventanas_p": {
        "title": "5. Climatología Mensual de Precipitación por Ventana Temporal",
        "files": [
            ("05_Climatologia_Mensual_Precipitacion/base_1981-2010.png", "Base (1981-2010)"),
            ("05_Climatologia_Mensual_Precipitacion/cercano_2021-2050.png", "Cercano (2021-2050)"),
            ("05_Climatologia_Mensual_Precipitacion/medio_2041-2070.png", "Medio (2041-2070)"),
            ("05_Climatologia_Mensual_Precipitacion/tardio_2071-2100.png", "Tardío (2071-2100)"),
        ]
    },
    "6_climatologia_mensual_ventanas_pet": {
        "title": "6. Climatología Mensual de Evapotranspiración por Ventana Temporal",
        "files": [
            ("06_Climatologia_Mensual_Evapotranspiracion/base_1981-2010.png", "Base (1981-2010)"),
            ("06_Climatologia_Mensual_Evapotranspiracion/cercano_2021-2050.png", "Cercano (2021-2050)"),
            ("06_Climatologia_Mensual_Evapotranspiracion/medio_2041-2070.png", "Medio (2041-2070)"),
            ("06_Climatologia_Mensual_Evapotranspiracion/tardio_2071-2100.png", "Tardío (2071-2100)"),
        ]
    },
    "7_climatologia_mensual_ventanas_wb": {
        "title": "7. Climatología Mensual de Balance Hídrico por Ventana Temporal",
        "files": [
            ("07_Climatologia_Mensual_Balance_Hidrico/base_1981-2010.png", "Base (1981-2010)"),
            ("07_Climatologia_Mensual_Balance_Hidrico/cercano_2021-2050.png", "Cercano (2021-2050)"),
            ("07_Climatologia_Mensual_Balance_Hidrico/medio_2041-2070.png", "Medio (2041-2070)"),
            ("07_Climatologia_Mensual_Balance_Hidrico/tardio_2071-2100.png", "Tardío (2071-2100)"),
        ]
    },
    "8_barras_por_ventana": {
        "title": "8. Valores Anuales por Ventana Temporal y Escenario",
        "files": [
            ("08_Barras_por_Ventana_Temporal/precipitacion_por_ventana.png", "Precipitación anual (mm/año)"),
            ("08_Barras_por_Ventana_Temporal/evapotranspiracion_por_ventana.png", "Evapotranspiración potencial anual (mm/año)"),
            ("08_Barras_por_Ventana_Temporal/balance_hidrico_por_ventana.png", "Balance hídrico anual (mm/año)"),
        ]
    },
    "9_cambios_balance_hidrico": {
        "title": "9. Cambios en Balance Hídrico por Escenario (Δ vs 1981-2010)",
        "files": [
            ("09_Cambios_Balance_Hidrico/delta_WB_ssp126.png", "Δ Balance Hídrico - SSP1-2.6"),
            ("09_Cambios_Balance_Hidrico/delta_WB_ssp370.png", "Δ Balance Hídrico - SSP3-7.0"),
            ("09_Cambios_Balance_Hidrico/delta_WB_ssp585.png", "Δ Balance Hídrico - SSP5-8.5"),
        ]
    },
    "10_mapas_componentes_base": {
        "title": "10. Mapas Espaciales - Período Base (1981-2010)",
        "files": [
            ("10_Mapas_Componentes_Base/climatologia_base_P_PET_WB.png", "Climatología base: Precipitación, PET y Balance Hídrico (mm/año)"),
        ]
    },
    "11_mapas_cambios_precipitacion": {
        "title": "11. Mapas de Cambios en Precipitación (Δ vs Base)",
        "files": [
            ("11_Mapas_Cambios_Precipitacion/delta_P_cercano_2021-2040.png", "Δ Precipitación 2021-2040 (SSP1-2.6, SSP3-7.0, SSP5-8.5)"),
            ("11_Mapas_Cambios_Precipitacion/delta_P_medio_2041-2070.png", "Δ Precipitación 2041-2070"),
            ("11_Mapas_Cambios_Precipitacion/delta_P_tardio_2071-2100.png", "Δ Precipitación 2071-2100"),
        ]
    },
    "12_mapas_cambios_pet": {
        "title": "12. Mapas de Cambios en Evapotranspiración (Δ vs Base)",
        "files": [
            ("12_Mapas_Cambios_Evapotranspiracion/delta_PET_cercano_2021-2040.png", "Δ Evapotranspiración 2021-2040"),
            ("12_Mapas_Cambios_Evapotranspiracion/delta_PET_medio_2041-2070.png", "Δ Evapotranspiración 2041-2070"),
            ("12_Mapas_Cambios_Evapotranspiracion/delta_PET_tardio_2071-2100.png", "Δ Evapotranspiración 2071-2100"),
        ]
    },
    "13_mapas_cambios_balance": {
        "title": "13. Mapas de Cambios en Balance Hídrico (Δ vs Base)",
        "files": [
            ("13_Mapas_Cambios_Balance_Hidrico/delta_WB_cercano_2021-2040.png", "Δ Balance Hídrico 2021-2040"),
            ("13_Mapas_Cambios_Balance_Hidrico/delta_WB_medio_2041-2070.png", "Δ Balance Hídrico 2041-2070"),
            ("13_Mapas_Cambios_Balance_Hidrico/delta_WB_tardio_2071-2100.png", "Δ Balance Hídrico 2071-2100"),
        ]
    },
    "14_mapas_multiventana": {
        "title": "14. Matriz de Balance Hídrico Multi-Ventana",
        "files": [
            ("14_Matriz_MultiVentana/matriz_WB_escenarios_ventanas.png", "Balance Hídrico por Escenario y Ventana Temporal (Base y Δ)"),
        ]
    },
    "15_trimestres_base": {
        "title": "15. Análisis Estacional - Trimestres Húmedo y Seco (Base)",
        "files": [
            ("15_Trimestres_Base/WB_trimestres_base_1981-2010.png", "Balance Hídrico Trimestral - Período Base 1981-2010"),
        ]
    },
    "16_trimestres_cambios": {
        "title": "16. Cambios en Trimestres por Escenario",
        "files": [
            ("16_Trimestres_Cambios/delta_trimestres_ssp126.png", "Δ Balance Hídrico Trimestral - SSP1-2.6"),
            ("16_Trimestres_Cambios/delta_trimestres_ssp370.png", "Δ Balance Hídrico Trimestral - SSP3-7.0"),
            ("16_Trimestres_Cambios/delta_trimestres_ssp585.png", "Δ Balance Hídrico Trimestral - SSP5-8.5"),
        ]
    },
    "17_mapas_mensuales_historico": {
        "title": "17. Mapas Mensuales de Balance Hídrico - Período Base",
        "files": [
            ("17_Mapas_Mensuales_Historico/WB_mensual_historico_base.png", "Balance Hídrico Mensual Base (12 meses)"),
        ]
    },
    "18_mapas_mensuales_ssp126": {
        "title": "18. Mapas Mensuales de Balance Hídrico - SSP1-2.6",
        "files": [
            ("18_Mapas_Mensuales_SSP126/WB_mensual_ssp126_cercano.png", "Cercano (2021-2040)"),
            ("18_Mapas_Mensuales_SSP126/WB_mensual_ssp126_medio.png", "Medio (2041-2070)"),
            ("18_Mapas_Mensuales_SSP126/WB_mensual_ssp126_tardio.png", "Tardío (2071-2100)"),
        ]
    },
    "19_mapas_mensuales_ssp370": {
        "title": "19. Mapas Mensuales de Balance Hídrico - SSP3-7.0",
        "files": [
            ("19_Mapas_Mensuales_SSP370/WB_mensual_ssp370_cercano.png", "Cercano (2021-2040)"),
            ("19_Mapas_Mensuales_SSP370/WB_mensual_ssp370_medio.png", "Medio (2041-2070)"),
            ("19_Mapas_Mensuales_SSP370/WB_mensual_ssp370_tardio.png", "Tardío (2071-2100)"),
        ]
    },
    "20_mapas_mensuales_ssp585": {
        "title": "20. Mapas Mensuales de Balance Hídrico - SSP5-8.5",
        "files": [
            ("20_Mapas_Mensuales_SSP585/WB_mensual_ssp585_cercano.png", "Cercano (2021-2040)"),
            ("20_Mapas_Mensuales_SSP585/WB_mensual_ssp585_medio.png", "Medio (2041-2070)"),
            ("20_Mapas_Mensuales_SSP585/WB_mensual_ssp585_tardio.png", "Tardío (2071-2100)"),
        ]
    },
    "21_mapas_mensuales_delta_ssp126": {
        "title": "21. Mapas Mensuales de Cambios en Balance Hídrico - SSP1-2.6 (Δ vs Base)",
        "files": [
            ("21_Mapas_Mensuales_Delta_SSP126/delta_WB_mensual_ssp126_cercano.png", "Cercano (2021-2040)"),
            ("21_Mapas_Mensuales_Delta_SSP126/delta_WB_mensual_ssp126_medio.png", "Medio (2041-2070)"),
            ("21_Mapas_Mensuales_Delta_SSP126/delta_WB_mensual_ssp126_tardio.png", "Tardío (2071-2100)"),
        ]
    },
    "22_mapas_mensuales_delta_ssp370": {
        "title": "22. Mapas Mensuales de Cambios en Balance Hídrico - SSP3-7.0 (Δ vs Base)",
        "files": [
            ("22_Mapas_Mensuales_Delta_SSP370/delta_WB_mensual_ssp370_cercano.png", "Cercano (2021-2040)"),
            ("22_Mapas_Mensuales_Delta_SSP370/delta_WB_mensual_ssp370_medio.png", "Medio (2041-2070)"),
            ("22_Mapas_Mensuales_Delta_SSP370/delta_WB_mensual_ssp370_tardio.png", "Tardío (2071-2100)"),
        ]
    },
    "23_mapas_mensuales_delta_ssp585": {
        "title": "23. Mapas Mensuales de Cambios en Balance Hídrico - SSP5-8.5 (Δ vs Base)",
        "files": [
            ("23_Mapas_Mensuales_Delta_SSP585/delta_WB_mensual_ssp585_cercano.png", "Cercano (2021-2040)"),
            ("23_Mapas_Mensuales_Delta_SSP585/delta_WB_mensual_ssp585_medio.png", "Medio (2041-2070)"),
            ("23_Mapas_Mensuales_Delta_SSP585/delta_WB_mensual_ssp585_tardio.png", "Tardío (2071-2100)"),
        ]
    },
}


def add_title_page(doc, regions):
    """Añade página de portada."""

    title = doc.add_heading("ANEXO DE FIGURAS", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)


    subtitle = doc.add_heading("Análisis de Cambio Climático", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(20)


    subtitle2 = doc.add_heading("Proyecciones Hidrológicas y Climáticas", 1)
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2_run = subtitle2.runs[0]
    subtitle2_run.font.size = Pt(16)

    doc.add_paragraph()
    doc.add_paragraph()


    regions_para = doc.add_paragraph()
    regions_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    region_names = [info["name"] for info in regions.values()]
    regions_text = "Regiones Analizadas:\n\n" + "\n".join(region_names)
    regions_run = regions_para.add_run(regions_text)
    regions_run.font.size = Pt(14)
    regions_run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()


    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_run = period_para.add_run(f"Período de Análisis: {settings.PERIOD_START}-{settings.PERIOD_END}\nPeríodo Base: {settings.BASE_PERIOD[0]}-{settings.BASE_PERIOD[1]}")
    period_run.font.size = Pt(12)


    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Fecha de generación: {datetime.now().strftime('%d de %B de %Y')}")
    date_run.font.size = Pt(11)
    date_run.font.italic = True


    doc.add_page_break()


def add_toc_placeholder(doc):
    """Añade placeholder para tabla de contenidos."""
    toc_heading = doc.add_heading("ÍNDICE", 0)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("(Tabla de contenidos - generar en Word con F9 o actualizar campos)")
    doc.add_paragraph()
    doc.add_page_break()


def add_section_heading(doc, region_name):
    """Añade encabezado de sección regional."""
    heading = doc.add_heading(f"REGIÓN: {region_name.upper()}", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.runs[0]
    heading_run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_page_break()


def add_figure(doc, image_path, caption, max_width=6.5):
    """Añade una figura con pie de foto."""
    if os.path.exists(image_path):
        try:

            doc.add_picture(image_path, width=Inches(max_width))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(f"Figura: {caption}")
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(64, 64, 64)

            return True
        except Exception as e:
            print(f"  ⚠ Error añadiendo imagen {image_path}: {e}")
            return False
    else:
        print(f"  ⚠ Archivo no encontrado: {image_path}")
        return False


def create_document(specific_regions=None, report_dir=None, regions=None):
    """Crea el documento completo.

    Args:
        specific_regions (list): Lista de códigos de región para procesar. Si es None, procesa todas.
        report_dir (str): Carpeta de salida del .docx. Si es None usa settings.REPORTS_DIR.
        regions (dict): Diccionario de regiones a considerar. Si es None usa settings.REGIONS.
    """
    print("Creando documento Word con figuras...")

    active_regions = regions or settings.REGIONS
    output_reports_dir = report_dir or settings.REPORTS_DIR

    os.makedirs(output_reports_dir, exist_ok=True)


    doc = Document()


    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


    print("\n📄 Generando portada...")
    add_title_page(doc, active_regions)


    print("📑 Añadiendo índice...")
    add_toc_placeholder(doc)


    for region_code, region_info in active_regions.items():
        if specific_regions and region_code not in specific_regions:
            continue

        region_name = region_info["name"]
        region_dir = settings.get_region_output_dir(region_code)

        print(f"\n{'='*60}")
        print(f"🌍 Procesando región: {region_name}")
        print(f"{'='*60}")


        add_section_heading(doc, region_name)


        total_figs = 0
        added_figs = 0


        for category_key, category_info in FILE_STRUCTURE.items():
            category_title = category_info["title"]
            files_list = category_info["files"]

            print(f"\n  📂 {category_title}")


            doc.add_heading(category_title, 2)


            for file_path, caption in files_list:
                full_path = os.path.join(region_dir, file_path)
                total_figs += 1

                if add_figure(doc, full_path, caption):
                    added_figs += 1
                    print(f"    ✅ {caption}")
                    doc.add_paragraph()
                else:
                    print(f"    ❌ {caption}")


            doc.add_page_break()

        print(f"\n  📊 Resumen {region_name}: {added_figs}/{total_figs} figuras añadidas")


    output_path = os.path.join(output_reports_dir, f"Anexo_Figuras_Cambio_Climatico_{datetime.now().strftime('%Y%m%d')}.docx")
    doc.save(output_path)

    print(f"\n{'='*60}")
    print(f"✅ Documento creado exitosamente:")
    print(f"   {output_path}")
    print(f"{'='*60}")

    return output_path


if __name__ == "__main__":
    try:
        output_file = create_document()
        print("\n🎉 ¡Proceso completado!")
        print("\n📝 Instrucciones:")
        print("   1. Abrir el documento en Microsoft Word")
        print("   2. Ir al índice y presionar F9 (o botón derecho → Actualizar campo)")
        print("   3. Revisar y ajustar formato según necesidad")
        print("\n")
    except Exception as e:
        print(f"\n❌ Error durante la creación del documento:")
        print(f"   {e}")
        import traceback
        traceback.print_exc()
